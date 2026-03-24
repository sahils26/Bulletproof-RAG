import pytest
import os
import docx
from pypdf import PdfWriter

from deeprag.loaders import (
    PlainTextLoader, MarkdownLoader, PDFLoader, DOCXLoader,
    LoaderRegistry, UnsupportedFileTypeError, default_registry,
)


@pytest.fixture
def sample_text_file(tmp_path):
    p = tmp_path / "sample.txt"
    p.write_text("This is a simple text file.\nWe can read it.", encoding="utf-8")
    return str(p)


@pytest.fixture
def sample_md_file(tmp_path):
    p = tmp_path / "sample.md"
    p.write_text("# Main Header\nSome intro text.\n## Sub Header\nMore details.", encoding="utf-8")
    return str(p)


@pytest.fixture
def sample_docx_file(tmp_path):
    p = str(tmp_path / "sample.docx")
    doc = docx.Document()
    doc.add_paragraph("First paragraph in Word.")
    doc.add_paragraph("Second paragraph in Word.")
    doc.save(p)
    return p


@pytest.fixture
def sample_pdf_file(tmp_path):
    p = str(tmp_path / "sample.pdf")
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with open(p, "wb") as f:
        writer.write(f)
    return p


def test_plaintext_loader(sample_text_file):
    loader = PlainTextLoader()
    docs = loader.load(sample_text_file)
    assert len(docs) == 1
    assert "simple text file" in docs[0].content
    assert docs[0].metadata["file_type"] == "text/plain"


def test_markdown_loader(sample_md_file):
    loader = MarkdownLoader()
    docs = loader.load(sample_md_file)
    assert len(docs) == 1
    assert "Sub Header" in docs[0].content
    assert "Main Header" in docs[0].metadata["headers"]
    assert "Sub Header" in docs[0].metadata["headers"]


def test_docx_loader(sample_docx_file):
    loader = DOCXLoader()
    docs = loader.load(sample_docx_file)
    assert len(docs) == 1
    assert "First paragraph in Word." in docs[0].content
    assert "Second paragraph in Word." in docs[0].content


def test_pdf_loader(sample_pdf_file):
    loader = PDFLoader()
    docs = loader.load(sample_pdf_file)
    assert len(docs) == 1
    assert docs[0].metadata["total_pages"] == 1
    assert docs[0].content == ""


def test_supported_extensions():
    assert ".txt" in PlainTextLoader().supported_extensions()
    assert ".md" in MarkdownLoader().supported_extensions()
    assert ".pdf" in PDFLoader().supported_extensions()
    assert ".docx" in DOCXLoader().supported_extensions()


def test_loader_registry_dispatch(sample_text_file, sample_md_file):
    registry = default_registry()
    docs_txt = registry.load(sample_text_file)
    assert len(docs_txt) == 1
    assert docs_txt[0].metadata["file_type"] == "text/plain"

    docs_md = registry.load(sample_md_file)
    assert len(docs_md) == 1
    assert docs_md[0].metadata["file_type"] == "text/markdown"


def test_loader_registry_unknown_extension(tmp_path):
    registry = default_registry()
    fake_file = str(tmp_path / "file.xyz")
    with open(fake_file, "w") as f:
        f.write("data")
    with pytest.raises(UnsupportedFileTypeError):
        registry.load(fake_file)
