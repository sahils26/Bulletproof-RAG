import os
from abc import ABC, abstractmethod
from typing import List
from pathlib import Path
import re

from shared.models import Document


class UnsupportedFileTypeError(Exception):
    """Raised when no loader is registered for a given file extension."""
    pass


class DocumentLoader(ABC):
    """Abstract interface for all document loaders."""

    @abstractmethod
    def load(self, path: str) -> List[Document]:
        """Load text from the given path and return a list of Documents."""
        pass

    @abstractmethod
    def supported_extensions(self) -> List[str]:
        """Return list of file extensions this loader handles (e.g. ['.txt'])."""
        pass


class PlainTextLoader(DocumentLoader):
    """Loads a simple `.txt` file."""

    def supported_extensions(self) -> List[str]:
        return [".txt"]

    def load(self, path: str) -> List[Document]:
        if not os.path.exists(path):
            raise FileNotFoundError(f"File not found: {path}")

        with open(path, "r", encoding="utf-8") as f:
            text = f.read()

        doc = Document(
            source_path=path,
            content=text,
            metadata={
                "source": path,
                "file_type": "text/plain",
                "filename": os.path.basename(path),
            },
        )
        return [doc]


class MarkdownLoader(DocumentLoader):
    """Loads a `.md` file and extracts top-level headers into metadata."""

    def supported_extensions(self) -> List[str]:
        return [".md", ".markdown"]

    def load(self, path: str) -> List[Document]:
        if not os.path.exists(path):
            raise FileNotFoundError(f"File not found: {path}")

        with open(path, "r", encoding="utf-8") as f:
            text = f.read()

        headers = re.findall(r"^(#{1,6})\s+(.+)$", text, re.MULTILINE)
        extracted_headers = [h[1].strip() for h in headers]

        doc = Document(
            source_path=path,
            content=text,
            metadata={
                "source": path,
                "file_type": "text/markdown",
                "filename": os.path.basename(path),
                "headers": extracted_headers,
            },
        )
        return [doc]


class PDFLoader(DocumentLoader):
    """Loads text from `.pdf` files using pypdf."""

    def supported_extensions(self) -> List[str]:
        return [".pdf"]

    def load(self, path: str) -> List[Document]:
        import pypdf

        if not os.path.exists(path):
            raise FileNotFoundError(f"File not found: {path}")

        text_pages = []
        try:
            with open(path, "rb") as f:
                reader = pypdf.PdfReader(f)
                num_pages = len(reader.pages)
                for page_num in range(num_pages):
                    page = reader.pages[page_num]
                    extracted = page.extract_text()
                    if extracted:
                        text_pages.append(extracted)

            full_text = "\n\n".join(text_pages)

            doc = Document(
                source_path=path,
                content=full_text,
                metadata={
                    "source": path,
                    "file_type": "application/pdf",
                    "filename": os.path.basename(path),
                    "total_pages": num_pages,
                    "ocr_required": len(text_pages) == 0,
                },
            )
            return [doc]
        except Exception as e:
            raise RuntimeError(f"Failed to load PDF {path}: {str(e)}")


class DOCXLoader(DocumentLoader):
    """Loads text from `.docx` files using python-docx."""

    def supported_extensions(self) -> List[str]:
        return [".docx"]

    def load(self, path: str) -> List[Document]:
        import docx

        if not os.path.exists(path):
            raise FileNotFoundError(f"File not found: {path}")

        try:
            doc_file = docx.Document(path)
            full_text = "\n".join(
                [para.text for para in doc_file.paragraphs if para.text.strip()]
            )

            doc = Document(
                source_path=path,
                content=full_text,
                metadata={
                    "source": path,
                    "file_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "filename": os.path.basename(path),
                },
            )
            return [doc]
        except Exception as e:
            raise RuntimeError(f"Failed to load DOCX {path}: {str(e)}")


class LoaderRegistry:
    """Auto-dispatches to the correct loader based on file extension."""

    def __init__(self) -> None:
        self._loaders: dict[str, DocumentLoader] = {}

    def register(self, loader: DocumentLoader) -> None:
        """Register a loader for all of its declared extensions."""
        for ext in loader.supported_extensions():
            self._loaders[ext.lower()] = loader

    def load(self, path: str) -> List[Document]:
        """Load a file by dispatching to the correct loader based on extension."""
        ext = os.path.splitext(path)[1].lower()
        if ext not in self._loaders:
            raise UnsupportedFileTypeError(
                f"No loader registered for extension '{ext}'. "
                f"Supported: {list(self._loaders.keys())}"
            )
        return self._loaders[ext].load(path)


def default_registry() -> LoaderRegistry:
    """Create a registry pre-loaded with all built-in loaders."""
    registry = LoaderRegistry()
    registry.register(PlainTextLoader())
    registry.register(MarkdownLoader())
    registry.register(PDFLoader())
    registry.register(DOCXLoader())
    return registry
